"""Conversation-scoped export link and file rendering helpers."""

from __future__ import annotations

import base64
import csv
import hashlib
import hmac
import html
import io
import json
import mimetypes
import os
import re
import time
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Awaitable, Callable, cast
from urllib.parse import quote

from docx import Document
from fastapi import HTTPException
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

from ragtime.config import settings
from ragtime.core.logging import get_logger

logger = get_logger(__name__)

EXPORT_SPEC_VERSION = 1
EXPORT_DEFAULT_TTL_SECONDS = 60 * 60
EXPORT_MAX_TTL_SECONDS = 24 * 60 * 60
EXPORT_MAX_ROWS = 100_000
EXPORT_MAX_CELL_CHARS = 32_000
EXPORT_MAX_CONTENT_BYTES = 25 * 1024 * 1024
EXPORT_BASE_DIR = Path(settings.index_data_path) / "_exports" / "conversation_downloads"

SUPPORTED_GENERATED_FORMATS = {
    "csv",
    "tsv",
    "xlsx",
    "xls",
    "json",
    "txt",
    "text",
    "md",
    "markdown",
    "html",
    "xml",
    "pdf",
    "doc",
    "docx",
}

MIME_TYPES: dict[str, str] = {
    "csv": "text/csv; charset=utf-8",
    "tsv": "text/tab-separated-values; charset=utf-8",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "json": "application/json; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
    "text": "text/plain; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "markdown": "text/markdown; charset=utf-8",
    "html": "text/html; charset=utf-8",
    "xml": "application/xml; charset=utf-8",
    "pdf": "application/pdf",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

LiveTableResolver = Callable[[dict[str, Any]], Awaitable[tuple[list[str], list[list[Any]]]]]


def utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _parse_utc(value: str) -> datetime:
    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _secret() -> bytes:
    secret = (settings.encryption_key or settings.api_key or "ragtime-export-links").encode("utf-8")
    return hashlib.sha256(secret).digest()


def _base64_url_encode(payload: bytes) -> str:
    return base64.urlsafe_b64encode(payload).decode("ascii").rstrip("=")


def _base64_url_decode(payload: str) -> bytes:
    padding = "=" * (-len(payload) % 4)
    return base64.urlsafe_b64decode((payload + padding).encode("ascii"))


def normalize_export_format(value: str) -> str:
    fmt = re.sub(r"[^a-zA-Z0-9]+", "", (value or "").strip().lower())
    if fmt == "text":
        return "txt"
    if fmt == "markdown":
        return "md"
    return fmt or "txt"


def sanitize_export_filename(filename: str, export_format: str) -> str:
    export_format = normalize_export_format(export_format)
    raw = (filename or "export").strip().replace("\\", "/").split("/")[-1]
    raw = re.sub(r"[\x00-\x1f\x7f]", "", raw)
    raw = re.sub(r"\s+", "_", raw)
    raw = re.sub(r"[^A-Za-z0-9._ -]", "_", raw).strip(". _-") or "export"
    if "." in raw:
        stem, _, ext = raw.rpartition(".")
        if normalize_export_format(ext) != export_format:
            raw = f"{stem or 'export'}.{export_format}"
    else:
        raw = f"{raw}.{export_format}"
    return raw[:180]


def _safe_conversation_dir(conversation_id: str) -> Path:
    safe = re.sub(r"[^A-Za-z0-9_.-]", "_", conversation_id or "unknown")[:128]
    path = EXPORT_BASE_DIR / safe
    path.mkdir(parents=True, exist_ok=True)
    return path


def _spec_path(conversation_id: str, export_id: str) -> Path:
    safe_export_id = re.sub(r"[^A-Za-z0-9_-]", "", export_id or "")
    if not safe_export_id:
        raise HTTPException(status_code=400, detail="Invalid export id")
    base = _safe_conversation_dir(conversation_id).resolve()
    path = (base / f"{safe_export_id}.json").resolve()
    if not str(path).startswith(str(base)):
        raise HTTPException(status_code=400, detail="Invalid export path")
    return path


def _coerce_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return json.dumps(value, ensure_ascii=False, default=str)


def _as_dict(value: Any) -> dict[str, Any]:
    return cast(dict[str, Any], value) if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return cast(list[Any], value) if isinstance(value, list) else []


def _coerce_cell(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, (int, float, bool)):
        return value
    if isinstance(value, (dict, list, tuple)):
        value = json.dumps(value, ensure_ascii=False, default=str)
    else:
        value = str(value)
    return value[:EXPORT_MAX_CELL_CHARS] if len(value) > EXPORT_MAX_CELL_CHARS else value


def normalize_table(columns: Any, rows: Any) -> tuple[list[str], list[list[Any]]]:
    normalized_rows: list[list[Any]] = []
    if isinstance(rows, list):
        for row in rows[:EXPORT_MAX_ROWS]:
            if isinstance(row, dict):
                normalized_rows.append([_coerce_cell(value) for value in row.values()])
            elif isinstance(row, (list, tuple)):
                normalized_rows.append([_coerce_cell(value) for value in row])
            else:
                normalized_rows.append([_coerce_cell(row)])

    normalized_columns: list[str] = []
    if isinstance(columns, list):
        for index, column in enumerate(columns):
            if isinstance(column, dict):
                title = column.get("title") or column.get("name") or column.get("data") or f"Column {index + 1}"
            else:
                title = column
            normalized_columns.append(str(title or f"Column {index + 1}"))

    if not normalized_columns:
        normalized_columns = [f"Column {index + 1}" for index in range(max((len(row) for row in normalized_rows), default=0))]

    width = len(normalized_columns)
    for row in normalized_rows:
        if len(row) < width:
            row.extend([""] * (width - len(row)))
        elif len(row) > width:
            del row[width:]
    return normalized_columns, normalized_rows


def chart_to_table(payload: dict[str, Any]) -> tuple[list[str], list[list[Any]]]:
    config = _as_dict(payload.get("config"))
    data = _as_dict(config.get("data"))
    labels = _as_list(data.get("labels"))
    datasets = _as_list(data.get("datasets"))
    columns = ["Label"]
    for index, dataset in enumerate(datasets):
        columns.append(str(dataset.get("label") or f"Dataset {index + 1}") if isinstance(dataset, dict) else f"Dataset {index + 1}")
    rows: list[list[Any]] = []
    for row_index, label in enumerate(labels):
        row = [label]
        for dataset in datasets:
            values = dataset.get("data") if isinstance(dataset, dict) else []
            row.append(values[row_index] if isinstance(values, list) and row_index < len(values) else "")
        rows.append(row)
    return normalize_table(columns, rows)


def datatable_to_table(payload: dict[str, Any]) -> tuple[list[str], list[list[Any]]]:
    config = _as_dict(payload.get("config"))
    return normalize_table(config.get("columns"), config.get("data"))


def create_token(conversation_id: str, export_id: str, filename: str, expires_at: datetime) -> str:
    payload = {"conversation_id": conversation_id, "export_id": export_id, "filename": filename, "exp": int(expires_at.timestamp())}
    payload_b64 = _base64_url_encode(json.dumps(payload, sort_keys=True, separators=(",", ":")).encode("utf-8"))
    signature = hmac.new(_secret(), payload_b64.encode("ascii"), hashlib.sha256).digest()
    return f"{payload_b64}.{_base64_url_encode(signature)}"


def verify_token(token: str, conversation_id: str, export_id: str, filename: str) -> dict[str, Any]:
    try:
        payload_b64, signature_b64 = token.split(".", 1)
        expected_signature = hmac.new(_secret(), payload_b64.encode("ascii"), hashlib.sha256).digest()
        if not hmac.compare_digest(expected_signature, _base64_url_decode(signature_b64)):
            raise ValueError("bad signature")
        payload = json.loads(_base64_url_decode(payload_b64).decode("utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=403, detail="Invalid export link") from exc

    if payload.get("conversation_id") != conversation_id or payload.get("export_id") != export_id or payload.get("filename") != filename:
        raise HTTPException(status_code=403, detail="Invalid export link")
    if int(payload.get("exp") or 0) < int(time.time()):
        raise HTTPException(status_code=410, detail="Export link expired")
    return payload


def build_download_url(conversation_id: str, export_id: str, filename: str, token: str, workspace_id: str | None = None) -> str:
    url = f"/indexes/conversations/{quote(conversation_id)}/exports/{quote(export_id)}/{quote(filename)}?token={quote(token)}"
    if workspace_id:
        url += f"&workspace_id={quote(workspace_id)}"
    return url


def create_export_spec(
    *,
    conversation_id: str,
    filename: str,
    export_format: str,
    source: dict[str, Any],
    workspace_id: str | None = None,
    title: str | None = None,
    mime_type: str | None = None,
    expires_in_seconds: int = EXPORT_DEFAULT_TTL_SECONDS,
) -> dict[str, Any]:
    export_format = normalize_export_format(export_format)
    if export_format not in SUPPORTED_GENERATED_FORMATS and source.get("kind") != "binary_snapshot":
        raise HTTPException(status_code=400, detail=f"Unsupported generated export format: {export_format}")
    ttl = max(60, min(EXPORT_MAX_TTL_SECONDS, int(expires_in_seconds or EXPORT_DEFAULT_TTL_SECONDS)))
    created_at = utc_now()
    expires_at = created_at + timedelta(seconds=ttl)
    export_id = uuid.uuid4().hex
    safe_filename = sanitize_export_filename(filename or title or "export", export_format)
    spec = {
        "version": EXPORT_SPEC_VERSION,
        "id": export_id,
        "conversation_id": conversation_id,
        "workspace_id": workspace_id or None,
        "filename": safe_filename,
        "format": export_format,
        "title": title or safe_filename.rsplit(".", 1)[0],
        "mime_type": mime_type or MIME_TYPES.get(export_format) or mimetypes.guess_type(safe_filename)[0] or "application/octet-stream",
        "source": source,
        "created_at": created_at.isoformat(),
        "expires_at": expires_at.isoformat(),
    }
    path = _spec_path(conversation_id, export_id)
    path.write_text(json.dumps(spec, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    token = create_token(conversation_id, export_id, safe_filename, expires_at)
    spec["download_url"] = build_download_url(conversation_id, export_id, safe_filename, token, workspace_id=workspace_id)
    spec["token"] = token
    return spec


def load_export_spec(conversation_id: str, export_id: str) -> dict[str, Any]:
    path = _spec_path(conversation_id, export_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Export not found")
    try:
        spec = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Export spec is unreadable") from exc
    try:
        expires_at = _parse_utc(str(spec.get("expires_at") or ""))
    except Exception as exc:
        raise HTTPException(status_code=410, detail="Export link expired") from exc
    if expires_at < utc_now():
        raise HTTPException(status_code=410, detail="Export link expired")
    return spec


def cleanup_expired_exports(now: datetime | None = None) -> int:
    if not EXPORT_BASE_DIR.exists():
        return 0
    cutoff = now or utc_now()
    removed = 0
    for path in EXPORT_BASE_DIR.glob("*/*.json"):
        try:
            spec = json.loads(path.read_text(encoding="utf-8"))
            if _parse_utc(str(spec.get("expires_at") or "")) >= cutoff:
                continue
        except Exception:
            pass
        try:
            path.unlink()
            removed += 1
        except OSError:
            pass
    for directory in EXPORT_BASE_DIR.glob("*"):
        try:
            if directory.is_dir() and not any(directory.iterdir()):
                directory.rmdir()
        except OSError:
            pass
    return removed


def table_source(columns: Any, rows: Any) -> dict[str, Any]:
    normalized_columns, normalized_rows = normalize_table(columns, rows)
    return {"kind": "table_snapshot", "columns": normalized_columns, "rows": normalized_rows}


def content_source(text: Any = None, content_base64: str | None = None, mime_type: str | None = None) -> dict[str, Any]:
    if content_base64:
        try:
            content_bytes = base64.b64decode(content_base64, validate=True)
        except Exception as exc:
            raise HTTPException(status_code=400, detail="content_base64 is not valid base64") from exc
        if len(content_bytes) > EXPORT_MAX_CONTENT_BYTES:
            raise HTTPException(status_code=413, detail="Export content is too large")
        return {"kind": "binary_snapshot", "content_base64": content_base64, "mime_type": mime_type or "application/octet-stream"}
    encoded = _coerce_text(text).encode("utf-8")
    if len(encoded) > EXPORT_MAX_CONTENT_BYTES:
        raise HTTPException(status_code=413, detail="Export content is too large")
    return {"kind": "content_snapshot", "text": encoded.decode("utf-8", errors="replace"), "mime_type": mime_type}


def live_table_source(data_connection: dict[str, Any], columns: Any = None, rows: Any = None) -> dict[str, Any]:
    source: dict[str, Any] = {"kind": "live_table", "data_connection": data_connection}
    if rows is not None or columns is not None:
        snapshot_columns, snapshot_rows = normalize_table(columns, rows)
        source["snapshot_columns"] = snapshot_columns
        source["snapshot_rows"] = snapshot_rows
    return source


def _table_as_records(columns: list[str], rows: list[list[Any]]) -> list[dict[str, Any]]:
    return [{columns[index]: row[index] if index < len(row) else "" for index in range(len(columns))} for row in rows]


def _spreadsheet_safe(value: Any) -> Any:
    value = _coerce_cell(value)
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@")):
        return "\t" + value
    return value


def _csv_bytes(columns: list[str], rows: list[list[Any]], delimiter: str = ",") -> bytes:
    buffer = io.StringIO(newline="")
    writer = csv.writer(buffer, delimiter=delimiter)
    writer.writerow([_spreadsheet_safe(column) for column in columns])
    for row in rows:
        writer.writerow([_spreadsheet_safe(cell) for cell in row])
    return ("\ufeff" + buffer.getvalue()).encode("utf-8")


def _xlsx_bytes(columns: list[str], rows: list[list[Any]], title: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    if ws is None:
        raise RuntimeError("Workbook did not create an active worksheet")
    ws.title = re.sub(r"[^A-Za-z0-9 _-]", "", title or "Export")[:31] or "Export"
    ws.append(columns)
    for row in rows:
        ws.append([_coerce_cell(cell) for cell in row])
    for column_index, column_cells in enumerate(ws.columns, start=1):
        max_len = max((len(str(cell.value or "")) for cell in column_cells), default=0)
        ws.column_dimensions[get_column_letter(column_index)].width = min(max(max_len + 2, 10), 60)
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _html_table(columns: list[str], rows: list[list[Any]], title: str) -> str:
    head = "".join(f"<th>{html.escape(str(column))}</th>" for column in columns)
    body = "".join("<tr>" + "".join(f"<td>{html.escape(str(_coerce_cell(cell)))}</td>" for cell in row) + "</tr>" for row in rows)
    return (
        '<!doctype html><html><head><meta charset="utf-8"><title>'
        + html.escape(title or "Export")
        + "</title></head><body><h1>"
        + html.escape(title or "Export")
        + '</h1><table border="1" cellspacing="0" cellpadding="4"><thead><tr>'
        + head
        + "</tr></thead><tbody>"
        + body
        + "</tbody></table></body></html>"
    )


def _text_from_table(columns: list[str], rows: list[list[Any]], title: str) -> str:
    lines = [title or "Export", "", "\t".join(columns)]
    lines.extend("\t".join(str(_coerce_cell(cell)) for cell in row) for row in rows)
    return "\n".join(lines).strip() + "\n"


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _simple_pdf_bytes(text: str, title: str) -> bytes:
    lines: list[str] = []
    for raw_line in (title + "\n\n" + text).splitlines():
        line = raw_line[:110]
        while line:
            lines.append(line[:95])
            line = line[95:]
        if raw_line == "":
            lines.append("")
    pages = [lines[i : i + 42] for i in range(0, max(1, len(lines)), 42)] or [[]]
    objects: list[bytes] = []
    catalog_id = 1
    pages_id = 2
    font_id = 3
    page_ids: list[int] = []
    next_id = 4
    for page_lines in pages:
        content_id = next_id
        page_id = next_id + 1
        next_id += 2
        page_ids.append(page_id)
        stream_lines = ["BT", "/F1 10 Tf", "50 770 Td", "14 TL"]
        for line in page_lines:
            stream_lines.append(f"({_pdf_escape(line)}) Tj")
            stream_lines.append("T*")
        stream_lines.append("ET")
        stream = "\n".join(stream_lines).encode("latin-1", errors="replace")
        objects.append(f"{content_id} 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream\nendobj\n")
        objects.append(
            f"{page_id} 0 obj\n<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>\nendobj\n".encode(
                "ascii"
            )
        )
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    fixed_objects = [
        f"{catalog_id} 0 obj\n<< /Type /Catalog /Pages {pages_id} 0 R >>\nendobj\n".encode("ascii"),
        f"{pages_id} 0 obj\n<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>\nendobj\n".encode("ascii"),
        f"{font_id} 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n".encode("ascii"),
    ]
    all_objects = fixed_objects + objects
    output = io.BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for obj in sorted(all_objects, key=lambda item: int(item.split(b" ", 1)[0])):
        offsets.append(output.tell())
        output.write(obj)
    xref_offset = output.tell()
    output.write(f"xref\n0 {len(offsets)}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(f"trailer\n<< /Size {len(offsets)} /Root {catalog_id} 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    return output.getvalue()


def _docx_bytes(text: str, title: str, columns: list[str] | None = None, rows: list[list[Any]] | None = None) -> bytes:
    document = Document()
    if title:
        document.add_heading(title, level=1)
    if columns and rows is not None:
        table = document.add_table(rows=1, cols=len(columns))
        for index, column in enumerate(columns):
            table.rows[0].cells[index].text = str(column)
        for row in rows:
            cells = table.add_row().cells
            for index, cell in enumerate(row[: len(columns)]):
                cells[index].text = str(_coerce_cell(cell))
    else:
        for paragraph in text.split("\n"):
            document.add_paragraph(paragraph)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


async def resolve_source_table(spec: dict[str, Any], live_table_resolver: LiveTableResolver | None = None) -> tuple[list[str], list[list[Any]]]:
    source = _as_dict(spec.get("source"))
    kind = source.get("kind")
    if kind == "table_snapshot":
        return normalize_table(source.get("columns"), source.get("rows"))
    if kind == "chart_snapshot":
        return chart_to_table(_as_dict(source.get("payload")))
    if kind == "datatable_snapshot":
        return datatable_to_table(_as_dict(source.get("payload")))
    if kind == "live_table":
        if live_table_resolver is not None:
            try:
                return normalize_table(*await live_table_resolver(source))
            except HTTPException:
                if "snapshot_rows" not in source and "snapshot_columns" not in source:
                    raise
            except Exception:
                if "snapshot_rows" not in source and "snapshot_columns" not in source:
                    raise
        if "snapshot_rows" in source or "snapshot_columns" in source:
            return normalize_table(source.get("snapshot_columns"), source.get("snapshot_rows"))
        raise HTTPException(status_code=400, detail="Live export requires a data resolver")
    if kind == "content_snapshot":
        return normalize_table(["Content"], [[source.get("text") or ""]])
    raise HTTPException(status_code=400, detail="Export source is not table-shaped")


async def render_export(spec: dict[str, Any], live_table_resolver: LiveTableResolver | None = None) -> tuple[bytes, str]:
    export_format = normalize_export_format(str(spec.get("format") or "txt"))
    source = _as_dict(spec.get("source"))
    title = str(spec.get("title") or spec.get("filename") or "Export")
    if source.get("kind") == "binary_snapshot":
        data = base64.b64decode(str(source.get("content_base64") or ""), validate=True)
        return data, str(source.get("mime_type") or spec.get("mime_type") or "application/octet-stream")
    if source.get("kind") == "content_snapshot" and export_format not in {"csv", "tsv", "xlsx", "xls", "json"}:
        text = _coerce_text(source.get("text"))
        if export_format == "pdf":
            return _simple_pdf_bytes(text, title), MIME_TYPES["pdf"]
        if export_format == "docx":
            return _docx_bytes(text, title), MIME_TYPES["docx"]
        if export_format == "doc":
            content = f'<!doctype html><html><head><meta charset="utf-8"></head><body><pre>{html.escape(text)}</pre></body></html>'
            return content.encode("utf-8"), MIME_TYPES["doc"]
        return text.encode("utf-8"), str(source.get("mime_type") or spec.get("mime_type") or MIME_TYPES.get(export_format) or "text/plain; charset=utf-8")
    columns, rows = await resolve_source_table(spec, live_table_resolver)
    if export_format == "csv":
        return _csv_bytes(columns, rows), MIME_TYPES["csv"]
    if export_format == "tsv":
        return _csv_bytes(columns, rows, delimiter="\t"), MIME_TYPES["tsv"]
    if export_format == "xlsx":
        return _xlsx_bytes(columns, rows, title), MIME_TYPES["xlsx"]
    if export_format == "xls":
        return _html_table(columns, rows, title).encode("utf-8"), MIME_TYPES["xls"]
    if export_format == "json":
        return json.dumps(_table_as_records(columns, rows), ensure_ascii=False, indent=2, default=str).encode("utf-8"), MIME_TYPES["json"]
    if export_format in {"html", "doc"}:
        return _html_table(columns, rows, title).encode("utf-8"), MIME_TYPES[export_format]
    if export_format == "xml":
        records = _table_as_records(columns, rows)
        items = []
        for record in records:
            fields = "".join(f'<field name="{html.escape(str(key))}">{html.escape(str(value))}</field>' for key, value in record.items())
            items.append(f"<row>{fields}</row>")
        return ('<?xml version="1.0" encoding="UTF-8"?><rows>' + "".join(items) + "</rows>").encode("utf-8"), MIME_TYPES["xml"]
    text = _text_from_table(columns, rows, title)
    if export_format == "pdf":
        return _simple_pdf_bytes(text, title), MIME_TYPES["pdf"]
    if export_format == "docx":
        return _docx_bytes(text, title, columns, rows), MIME_TYPES["docx"]
    return text.encode("utf-8"), MIME_TYPES.get(export_format) or "text/plain; charset=utf-8"


def content_disposition(filename: str) -> str:
    ascii_name = re.sub(r"[^A-Za-z0-9._ -]", "_", filename) or "export"
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{quote(filename)}"


def write_export_bytes_to_temp(data: bytes, filename: str) -> Path:
    EXPORT_BASE_DIR.mkdir(parents=True, exist_ok=True)
    temp_name = f"stream-{uuid.uuid4().hex}-{sanitize_export_filename(filename, filename.rsplit('.', 1)[-1] if '.' in filename else 'bin')}"
    path = EXPORT_BASE_DIR / temp_name
    path.write_bytes(data)
    return path


def remove_file_quietly(path: Path) -> None:
    try:
        os.remove(path)
    except OSError:
        pass
